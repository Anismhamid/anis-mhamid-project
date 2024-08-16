#=======================================| |  ניהול עובדים  | |===========================================
import xlsxwriter
from ctypes import alignment
from tkinter import *
from tkinter import ttk, StringVar, messagebox
from tkinter import END
from tkinter.commondialog import Dialog
import webbrowser
import pymysql
from datetime import datetime
import customtkinter as ctk
from tkinter import filedialog
import openpyxl
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import requests
from tkinter import Frame
from docx.shared import RGBColor
import ttkbootstrap  as cttk
from ttkbootstrap.toast import ToastNotification
from ttkbootstrap.tooltip import ToolTip
from ttkbootstrap.constants import *
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
#-----------------------------------------------


mainFont = font='Heebo'

#----------- database conniction -------------
hostname1 = 'localhost'
porta = 10017
username1 ='root'
passwd1 = 'root'
database1='local'

themes = ['flatly', 'cyborg', 'litera', 'minty', 'lumen', 'sandstone', 
'yeti', 'pulse', 'united', 'morph', 'journal', 'darkly', 
'superhero', 'solar', 'cosmo', 'vapor', 'simplex', 'cerculean']




#--------- | colors | --------------
primary = '#2E4057'
secondary = '#adb5bd'
success = '#02b875'
info = '#17a2b8'
warning = '#F1D302'
danger = '#C1292E'
light = '#f8f9fa'
dark = '#343a40'
darkBlue ='#0077b6'



class MainScreen:
    def __init__(self):
        super().__init__()
        self.new_root = ctk.CTk()
        self.new_root.geometry('996x760+320+20')
        self.new_root.title('ניהול עוהדים, Develobed by °Anis Zkaria Mhamid°')
        self.new_root.config(bd=0)
        self.new_root.minsize(True,True)
        # self.new_root.configure(fg_color=primary)



        def fullscreen(event):
            self.new_root.attributes('-fullscreen', True)

            def exit_fullscreen(event):
                self.new_root.attributes('-fullscreen', False)
            
            self.new_root.bind('<Escape>', exit_fullscreen)
        self.new_root.bind('<F11>', fullscreen)

#=============================================================| |  variables   | |========================================================
        WORKERS_NAME_var = StringVar()
        WORKERS_PHONE_var = StringVar()
        WORKERS_ID_var = StringVar()
        WORKERS_COMPANY_NAME_var = StringVar()
        WORKERS_COMPANY_MANAGER_NAME_var = StringVar()
        WORKERS_WORK_ADDRESS_var = StringVar()
        WORKERS_DATE_var = StringVar()
        WORKERS_WAGE_var = StringVar()
        WORKERS_Taken_var = StringVar()
        WORKERS_Hours_Var = StringVar()
        search_var = StringVar()
        delete_var = StringVar()
        Row_Number_var = IntVar()
# #=====================================



#-----------== | פונקצית שקיפות |
        def allpha(value):
            global alpha
            alpha = float(value)
            alpha_list = [alpha]
            if alpha_list is not None:
                self.new_root.attributes('-alpha', alpha_list)
            else:
                pass
#----------------------------------------


# ============= | |  Header   | | ===================================

        Header = cttk.Frame(self.new_root,height=29)
        Header.pack(fill='x',pady=4)
        Header.pack_propagate(False)

        Header3 = cttk.Frame(self.new_root)
        Header3.pack(fill='both')

        # main label | ======
        def laBel(master,text,side='top'):
            welcome_to_my_app =cttk.Label(master=master,text=text,font=(mainFont,16,'bold'))
            welcome_to_my_app.pack(side=side)
        laBel(Header,'ניהול עובדים')

        # main Button  | | ==========
        def kaftor(text,command,master,side='top'):
            buttons = cttk.Button(master=master,default='active',text=text,command=command,cursor='hand2')
            buttons.pack(side=side,pady=5,ipady=8,fill=X)
#=====================================================================================================================




#================================== | |  תפריט   | | ===================================

# | תמונות תפריט |
#----------------------------------------------------
        def toggle_menu():
            point = 0.900
            window_height = self.new_root.winfo_height() - 200

            def toggle_menu_dis():
                tuggle_btn.configure( command=toggle_menu, text='תפריט')

                for i in range(200):
                    tuggle_menu_frame.place(x=-i - 120, y=90, height=window_height, width=20)
                    tuggle_menu_frame.update()
                    tuggle_menu_frame.after(int(point))

            tuggle_menu_frame = Frame(master=self.new_root,borderwidth=1)
            tuggle_menu_frame.place(x=0, y=90, height=window_height, width=200)

            # Toggle menu display logic
            if tuggle_menu_frame.winfo_x() == 0:
                toggle_menu_dis()
            else:
                toggle_menu_dis()
            laBel(tuggle_menu_frame, 'שקיפות' )


            tuggle_btn.configure( command=toggle_menu_dis, text='סגור')

            # Transparency Slider
            cale_tk = ctk.CTkSlider(tuggle_menu_frame, command=allpha, from_=0, to=1)

            for i in range(200):
                tuggle_menu_frame.place(x=i - 200, y=80, height=window_height, width=200)
                tuggle_menu_frame.update()
                tuggle_menu_frame.after(int(point))
            cale_tk.pack(ipadx=10,pady=200)

        # Menu Button
        tuggle_btn = cttk.Button(
            Header,
            width=10,
            text='תפריט',
            compound='left',
            command=toggle_menu
        )
        tuggle_btn.place(x=0, y=0)




#========================================================================================



#======================= | |  צץ תצוגה   | | =============================

        style =cttk.Style()
        style.theme_use(themes[0])
        def change_theme(event):
            selected_theme = combobox.get()
            style.theme_use(selected_theme)
            style.configure("Treeview",

                font=(mainFont, 10, 'normal'),
                rowheight=22)

        combobox = cttk.Combobox(Header, values=themes,width=10,
                        style="Custom.TCombobox",
                                        validate='all',
                                        state='readonly',
                                        font=(mainFont,10),
                                        justify='center')
        combobox.place(x=200)
        combobox.set(themes[0]) 

        # Bind combobox selection to theme change function
        combobox.bind('<<ComboboxSelected>>', change_theme)

        # Set initial theme
        style.theme_use(themes[6])


        style.configure("Treeview",
                        background='#D3D3D3',
                        foreground=dark,
                        fieldbackground='#D3D3D3',
                        font=(mainFont, 10, 'normal'),
                        rowheight=22)

        style.map('Treeview',
                background=[('selected', '#0078D7')])

        style.configure("Treeview.Heading",
                        font=('Arial', 9, 'bold'))


        # | פונקצית ייצור צץ להצגת הנתונים |
        def create_treeview(parent_frame):
            treeview_frame = ttk.Frame(parent_frame)
            treeview_frame.pack(padx=0, pady=0, fill=cttk.BOTH, expand=True)

            vertical_scrollbar = ttk.Scrollbar(treeview_frame, orient="vertical")
            vertical_scrollbar.pack(side='right', fill='y')


            workers_treeview = cttk.Treeview(
                treeview_frame,
                columns=('Rowid', "hours", 'taken', "wage", 'address', 'managername', 'compname', 'date', 'worker_id', 'phone', 'workername'),
                yscrollcommand=vertical_scrollbar.set,
                selectmode='extended', 
                height=9
            )
            workers_treeview.pack(fill=cttk.BOTH, expand=True)

            vertical_scrollbar.configure(command=workers_treeview.yview)

            workers_treeview.heading('Rowid', text="מ'ס שורה")
            workers_treeview.heading('hours', text="מ'ס שעות")
            workers_treeview.heading('taken', text='מפריעה')
            workers_treeview.heading('wage', text='שכר')
            workers_treeview.heading('address', text='כתובת')
            workers_treeview.heading('managername', text='מנהל')
            workers_treeview.heading('compname', text='חברה')
            workers_treeview.heading('date', text='תאריך')
            workers_treeview.heading('worker_id', text="ת'ז")
            workers_treeview.heading('phone', text='טלפון')
            workers_treeview.heading('workername', text='שם')

            workers_treeview.column('#0', width=0, stretch=cttk.NO, anchor=cttk.CENTER)
            workers_treeview.column('Rowid', width=40, anchor=cttk.CENTER)
            workers_treeview.column('hours', width=40, anchor=cttk.CENTER)
            workers_treeview.column('taken', width=50, anchor=cttk.CENTER)
            workers_treeview.column('wage', width=50, anchor=cttk.CENTER)
            workers_treeview.column('address', width=60, anchor=cttk.CENTER)
            workers_treeview.column('managername', width=80, anchor=cttk.CENTER)
            workers_treeview.column('compname', width=70, anchor=cttk.CENTER)
            workers_treeview.column('date', width=200, anchor=cttk.CENTER)
            workers_treeview.column('worker_id', width=80, anchor=cttk.CENTER)
            workers_treeview.column('phone', width=80, anchor=cttk.CENTER)
            workers_treeview.column('workername', width=70, anchor=cttk.CENTER)




            return workers_treeview
        workers_tuple = create_treeview(self.new_root)

#================ | |  Functions   | | ================================
        
# | פונקצית תוסט להוצאת תוצאות למשתמש 
        def toastErrorCacher(title,message,bg='light',position=(800,0,'ew')):
            toast = ToastNotification(title=title,bootstyle=bg,message=message,position=position,duration=10000,alert=True)
            toast.show_toast()

# | פונקצית הסתרת צץ הנתונים |
        def toggle_treeview():
                if workers_tuple.winfo_ismapped():
                    workers_tuple.pack_forget()
                else:
                    workers_tuple.pack(fill='both')


# | פונקצית קבלת הנתונים |==
        def get_api():
            try:
                # GET request to the API endpoint
                response = requests.get('http://127.0.0.1:5000/get')
                
                # Verify the success of the request
                if response.status_code == 200:
                    # Extract data into JSON format
                    data = response.json()
                    
                    # Clear existing treeview data
                    workers_tuple.delete(*workers_tuple.get_children())
                    count=0
                    # Inserting new data from the API into the Treeview
                    for row in data:
                        if count % 2==0:
                            workers_tuple.insert('', END, values=(
                                row['Rowid'],
                                row['hours'],
                                row['taken'],
                                row['wage'],
                                row['address'],
                                row['managername'],
                                row['compname'],
                                row['date'],
                                row['worker_id'],
                                row['phone'],
                                row['workername']
                            ), tags=('evenrow','odd'))
                else:
                    # If the request fails
                    toastErrorCacher('Error', f'Failed to fetch data. Status code: {response.status_code}')
                    print(response.status_code)
            
            except requests.exceptions.RequestException as e:
                print(e)
                toastErrorCacher('Error', f'An error occurred: {e}')
        get_api()


# | פונקצית שליחת הנתונים | 
        def post_api():

            data = {
                'WORKERS_NAME_var': WORKERS_NAME_var.get(),
                'WORKERS_PHONE_var': WORKERS_PHONE_var.get(),
                'WORKERS_ID_var': WORKERS_ID_var.get(),
                'WORKERS_COMPANY_NAME_var': WORKERS_COMPANY_NAME_var.get(),
                'WORKERS_COMPANY_MANAGER_NAME_var': WORKERS_COMPANY_MANAGER_NAME_var.get(),
                'WORKERS_WORK_ADDRESS_var': WORKERS_WORK_ADDRESS_var.get(),
                'WORKERS_DATE_var': WORKERS_DATE_var.get(),
                'WORKERS_WAGE_var': WORKERS_WAGE_var.get(),
                'WORKERS_Taken_var': WORKERS_Taken_var.get(),
                'WORKERS_Hours_Var': WORKERS_Hours_Var.get(),
            }

            headers = {'mainContent-Type': 'application/json'}

            url = 'http://127.0.0.1:5000/post'

            try:
                response = requests.post(url, json=data, headers=headers)

                if response.status_code == 201:
                    toastErrorCacher('הכנסת עובד','העובד הוקלט בהצלחה')
                else:

                    if response.status_code == 400:
                        toastErrorCacher('שגיאה', 'בקשה גרועה אנא בדוק את הנתונים שלך')
                    elif response.status_code == 500:
                        toastErrorCacher('שגיאה', 'שגיאת שרת פנימית בבקשה נסה שוב מאוחר יותר')
                    else:
                        toastErrorCacher('שגיאה', f'הוספת נתונים נכשלה. קוד סטטוס: {response.status_code}')

            except requests.exceptions.RequestException as e:
                # التعامل مع الأخطاء العامة في الطلبات (مثل مشاكل الشبكة)
                toastErrorCacher('שגיאה', f'אירעה שגיאה: {e}')
            except Exception as e:
                # التعامل مع أي استثناءات غير متوقعة
                toastErrorCacher('שגיאה', f'אירעה שגיאה בלתי צפויה: {e}')

# | פונקצית עדכון הנתונים |==
        def update_api():
            data = {
                'Rowid': Row_Number_var.get(),
                'hours': WORKERS_Hours_Var.get(),
                'taken': WORKERS_Taken_var.get(),
                'wage': WORKERS_WAGE_var.get(),
                'address': WORKERS_WORK_ADDRESS_var.get(),
                'managername': WORKERS_COMPANY_MANAGER_NAME_var.get(),
                'compname': WORKERS_COMPANY_NAME_var.get(),
                'date': WORKERS_DATE_var.get(),
                'worker_id': WORKERS_ID_var.get(),
                'phone': WORKERS_PHONE_var.get(),
                'workername': WORKERS_NAME_var.get(),
            }

            headers = {'mainContent-Type': 'application/json'}


            url = 'http://127.0.0.1:5000/update'
            

            try:
                # Send PUT request to the API endpoint
                response = requests.put(url, json=data,headers=headers)
                if response.status_code == 200:
                    toastErrorCacher('עדכון נתוני עובד','הנתונים של העובד עודכנו בהצלחה')
                else:
                    print(response)
                    # Handle specific error when updating data fails
                    if response.status_code == 400:
                        toastErrorCacher('שגיאה', 'בקשה גרועה אנא בדוק את הנתונים שלך')
                    elif response.status_code == 500:
                        toastErrorCacher('שגיאה', 'שגיאת שרת פנימית. בבקשה נסה שוב מאוחר יותר')
                    else:
                        toastErrorCacher('שגיאה', f'נכשל בעדכון הנתונים. קוד סטטוס: {response.status_code}')
            except requests.exceptions.RequestException as e:

                toastErrorCacher('שגיאה', f'אירעה שגיאה: {e}')
            except Exception as e:

                toastErrorCacher('שגיאה','אירעה שגיאה בלתי צפויה: {e}')


# | פונקצית מחיקת נתונים |==
        def delete_worker():
            try:
                worker_id = delete_var.get()
                response = requests.delete(f'http://127.0.0.1:5000/delete/{worker_id}')
                if response.status_code == 200:
                    toastErrorCacher('מחיקת עובד', 'הנתונים נמחקו בהצלחה','info')
                elif response.status_code == 404:
                    toastErrorCacher('שגיאה', 'העובד לא נמצא','danger')
                else:
                    toastErrorCacher('שגיאה', 'מחיקת הנתונים נכשלה. קוד סטטוס: {response.status_code}')
            except requests.exceptions.RequestException as e:
                toastErrorCacher('שגיאה', 'אירעה שגיאה: {e}','danger')
            except Exception as e:
                # Handle any other unexpected exceptions
                toastErrorCacher('שגיאה', f'אירעה שגיאה בלתי צפויה: {e}','danger')


# | פונקצית ניכוי כל שדות כל  |==
        def Clear_function():

            WORKERS_Hours_Var.set('')

            WORKERS_Taken_var.set('')
            
            WORKERS_WAGE_var.set('')

            WORKERS_WORK_ADDRESS_var.set('')

            WORKERS_COMPANY_MANAGER_NAME_var.set('')

            WORKERS_COMPANY_NAME_var.set('')

            WORKERS_ID_var.set('')

            WORKERS_PHONE_var.set('')

            WORKERS_NAME_var.set('')

            delete_var.set('')

            search_var.set('')


# | פונקצית קבלה | ניקוי שדות כל  |==
        def addandclear_function():
            post_api()
            Clear_function()
            get_api()
            get_workers_names()
            get_workers_names2(search2=search2)
            get_workers_names_Analysis()
            workers_tuple.xview_moveto(1)


# | פונקצית מחיקה | קבלה | ניקוי | שדות כל  |==
        def DeleteAndClear_function():
            delete_worker()
            get_api()
            get_workers_names()
            get_workers_names_Analysis()
            get_workers_names2(search2=search2)
            Clear_function()


# | ממלא את שדות הקלט כאשר אתה לוחץ על שם |
        def get_cursor_function(event):
            cursor_row = workers_tuple.focus()
            try:    
                if cursor_row:
                    mainContents = workers_tuple.item(cursor_row)
                    row = mainContents['values']
                    if row:
                        Row_Number_var.set(row[0])
                        WORKERS_Hours_Var.set(row[1])
                        WORKERS_Taken_var.set(row[2]),
                        WORKERS_WAGE_var.set(row[3])
                        WORKERS_WORK_ADDRESS_var.set(row[4])
                        WORKERS_COMPANY_MANAGER_NAME_var.set(row[5])
                        WORKERS_COMPANY_NAME_var.set(row[6])
                        WORKERS_DATE_var.set(row[7])
                        WORKERS_ID_var.set(row[8])
                        WORKERS_PHONE_var.set(row[9])
                        WORKERS_NAME_var.set(row[10])
            except pymysql.err.IntegrityError as e:
                toastErrorCacher('err', e,'danger')
            finally:
                pass
        workers_tuple.bind("<ButtonRelease-1>", get_cursor_function)


#| מחקה באמצעות כפתור delete |
        def delete_row_function(event):
            try:
                con = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                with con.cursor() as cor:
                    selected_items = workers_tuple.selection()
                    if selected_items:
                        for selected_item in selected_items:
                            row_id = workers_tuple.item(selected_item, 'values')[0]
                            confirmation = messagebox.askyesno("מוחק שורה", f"?האם אתה בטוח שברצונך למחוק שורה:_{row_id}")
                            if confirmation:
                                delete_query = "DELETE FROM workers WHERE `Rowid` = %s"
                                cor.execute(delete_query, (row_id,))
                                workers_tuple.delete(selected_item)
                        con.commit()
                        get_workers_names()
                        get_workers_names2(search2=search2)
                        dataAnalysis_func()
                    else:
                        toastErrorCacher("שגיאה", "לא נמחק",'danger')
            finally:
                pass
        workers_tuple.bind('<Delete>', delete_row_function)
        workers_tuple.bind('<Return>', post_api) 


        def validate_number(En1) -> bool:
            """Validates that the input is a number"""
            if En1.isdigit():
                return True
            elif En1 == "":
                return True
            else:
                return False

        def validate_alpha(En1) -> bool:
            """Validates that the input is alpha"""
            if En1.isdigit():
                return False
            elif En1 == "":
                return True
            else:
                return True


# | פונקצית עדכון נתונים במסד הנתונים | ניכוי שדות כלט  |
        def updateandclear():
            update_api()
            get_api()
            Clear_function()
            dataAnalysis_func()
            get_workers_names2(search2)
            get_workers_names_Analysis()
            workers_tuple.xview_moveto(1)


        def Search_function(event=None):
            try:
                searcha = search.get()
                con = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                cor = con.cursor()

                cor.execute("SELECT * FROM workers WHERE workername = %s", (searcha,))
                quary = cor.fetchall()

                if quary:
                    workers_tuple.delete(*workers_tuple.get_children())
                    for row in quary:
                        workers_tuple.insert('', END, values=row)
                    workers_tuple.xview_moveto(1.0)
                    con.commit()
                else:
                    toastErrorCacher("החיפוש שלך לא תאם","לא נמצאו עובדים התואמים לקריטריוני החיפוש שלך")
            except pymysql.Error as e:
                toastErrorCacher("שגיאה", "אירעה שגיאה בעת בדיקת מסד הנתונים")
            finally:
                con.close()


        #=== | פונקצית חישוב שכר | 
        def calculate_sum_function():
            try:
                conn = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                cursor = conn.cursor()
                sby = search2.get()
                if sby == 'כל העובדים':
                    query = f"SELECT SUM(`wage`) FROM workers"
                    cursor.execute(query)
                    res = cursor.fetchone()
                    total_wage = res[0]
                    toastErrorCacher('חישוב סך הכל',f'הסכום הכולל של כל העובדים: {total_wage}')
                    return total_wage
                else:
                    query = f"SELECT workername, SUM(`wage`) AS total_wage, SUM(Taken) AS total_taken FROM workers WHERE workername = '{sby}' GROUP BY workername"
                    cursor.execute(query)
                    res = cursor.fetchall()
                    if res:
                        for row in res:
                            workername, total_wage, total_taken = row
                            Min = total_wage - total_taken
                            result = messagebox.showinfo('חישוב סך הכל', f'חשבון סופי עובד : {workername} \n\n\n: {total_wage} \n\n\nקיבל :{total_taken} \n\n\n נטו : {Min}')
                            result
                            return total_wage, total_taken, Min
                    else:
                        toastErrorCacher("שגיאה","אין תוצאות, בחר שם עובד")
            except Exception as e:
                        toastErrorCacher('לא ניתן להתחבר לשרת MySQL ',f" לא ניתן היה ליצור חיבור מכיוון שמכונת היעד סירבה לכך באופן פעיל בדוק את חיבור האינטרנט שלך ונסה שנית")
            finally:
                return None


#-------------- | תאריך | ------------------ 
        def update_date():
            current_date = datetime.now().date()
            date_entry = Entry(Header3, font=(mainFont, 15, 'bold'),justify=CENTER, state='normal')
            date_entry.configure(bg=primary)
            date_entry.configure(fg=light)
            date_entry.insert(0, current_date)
            date_entry.pack(side='left',padx=0)
        update_date()


#--------------- | פונקצית שעון | ----------------
        def update_time():
            current_time = datetime.now()
            current_hour = current_time.hour
            current_minute = current_time.minute
            current_second = current_time.second
            formatted_time = "{} : {} : {}".format(current_hour, current_minute, current_second)
            date_entry2.delete(0, END)
            date_entry2.insert(0, formatted_time)
            Header3.after(1000, update_time) # |שדה שעון |
        date_entry2 = Entry(Header3,
                                justify=CENTER,
                                font=(mainFont, 15, 'bold'),
                                state='normal')
        date_entry2.configure(foreground=light)
        date_entry2.configure(background=primary)
        
        date_entry2.pack(side='right')
        update_time()


        # | פונקצית צור קשר |
        def infoo():
            toastErrorCacher('צור קשר' , 'Email: anesmhamed1@gmail.com\n\nPhone: 0538346915\n')
            url = "https://anismhamid.github.io/anis-mhamid-project/"
            webbrowser.open(url)


        # | kill program פונקצית|
        def des():
            self.new_root.destroy()



#======= | |  Printing Functions  | | =========================


        # Customize the table in DOCX
        def customize_table(table):
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(1, 2, 0)  # Black font color
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        customize_table(table)
        doc.save("output.docx")

        # Print to DOCX
        def print_to_docx():
            selected_value = search2.get()
            default_file_name = f'{selected_value}'
            file_path_docx = filedialog.asksaveasfilename(defaultextension=".docx",
                                                        initialfile=default_file_name,
                                                        filetypes=[("Word Files", "*.docx")])
            if file_path_docx:
                document = Document()
                title = document.add_heading(f"\n\n מידע עובדים\n", level=1)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table_data = []
                column_names = [workers_tuple.heading(column)["text"] for column in workers_tuple["columns"]]
                table_data.append(column_names)
                for item in workers_tuple.get_children():
                    item_values = [workers_tuple.item(item, "values")[i] for i in range(len(workers_tuple["columns"]))]
                    table_data.append(item_values)
                table = document.add_table(rows=len(table_data), cols=len(column_names))
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = str(cell_data)
                customize_table(table)
                document.save(file_path_docx)
                messagebox.showinfo('נשמר בהצלחה', f"{file_path_docx} :הקובץ נשמר ב")


        # Print to Excel
        def print_to_excel():
            selected_value = search.get()
            
            try:
                file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                        initialfile=selected_value,
                                                        filetypes=[("Excel Files", "*.xlsx")])
                
                if file_path:
                    workbook = openpyxl.Workbook()
                    sheet = workbook.active
                    sheet.sheet_view.rightToLeft = True
                    columns = workers_tuple["columns"]
                    column_headers = [workers_tuple.heading(column)["text"] for column in columns]
                    column_headers.reverse()
                    sheet.append(column_headers)
                    items = workers_tuple.get_children()

                    for row_idx, item in enumerate(items, start=2):
                        values = [workers_tuple.item(item, "values")[columns.index(column)] for column in columns]
                        values.reverse()

                        for col_idx, value in enumerate(values, start=1):
                            sheet.cell(row=row_idx, column=col_idx, value=value)

                    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row):
                        for cell in row:
                            cell.alignment = openpyxl.styles.Alignment(wrap_text=True)
                            cell.font = openpyxl.styles.Font(size=14)  # Font size can be adjusted
                            cell.alignment = openpyxl.styles.Alignment(vertical='center')  # Vertical alignment can be adjusted
                            sheet.row_dimensions[cell.row].height = 30

                    for col_idx, column_header in enumerate(column_headers, start=1):
                        col_letter = openpyxl.utils.get_column_letter(col_idx)
                        sheet.column_dimensions[col_letter].width = 20


                    for col_idx, column_header in enumerate(column_headers, start=1):
                        sheet.cell(row=1, column=col_idx, value=column_header)
                        sheet.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = len(column_header) +8
                    workbook.save(file_path)

                    if workbook:
                        messagebox.showinfo('הצלחה', f"הקובץ נשמר ב: {file_path}")
            except FileNotFoundError:
                toastErrorCacher('שגיאה', 'הקובץ לא נמצא.')
            except PermissionError:
                toastErrorCacher('שגיאה', 'הקובץ לא ניתן לכתיבה. נא לסגור אותו ולנסות שוב.')
            except Exception as e:
                print(e)
                toastErrorCacher('שגיאה', f'אירעה שגיאה: {str(e)}')




# ==================================================================



#================== | כפתורים עליונים | =================

        def topButtons(text,command,side):
            _top = cttk.Button(Header3,
                                width=150,
                                text=text,
                                command=command)
            _top.pack(side=side,padx='20')
        topButtons(text='הסתרת רשימה',command=toggle_treeview,side=RIGHT)



#===================================================



#=======| |  Calclate   | |=========

#                         | פרים חישוב | חיצוני  |
        fram_heshov = Frame(self.new_root)
        
        fram_heshov.pack(anchor='w',fill='x',pady=0,ipady=0)
#----------------------------------------------------

#                         | פרים חישוב פנימי  |
        fram_heshov2 = Frame(fram_heshov)
        fram_heshov2.configure(bg=primary)
        fram_heshov2.pack(anchor='s',fill='x',pady=(0,0))
#----------------------------------------------------

#                        | בחירת שם לחישוב |
        search2 = ttk.Combobox(fram_heshov2, state='readonly',width=15, font=( mainFont,12 ), justify='center')
        search2.pack(side='left',padx=(50,0),pady=(0,10))
#----------------------------------------------------

#                 | פונקצית ייבואי שימות מדאטאביס 2 |
        def get_workers_names2(search2):
            try:
                connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                cursor = connection.cursor()
                cursor.execute("SELECT DISTINCT workername FROM workers ORDER BY workername ASC;")
                rows = cursor.fetchall()
                global worker_names
                worker_names = [row[0] for row in rows]
                search2['values'] =['כל העובדים']+ worker_names
            except:
                pass
        get_workers_names2(search2)
#---------------------------------------------------------


#                         | כפתור חישוב |                
        calcbtn = cttk.Button(master = fram_heshov2,
                                text='חישוב',
                                command=calculate_sum_function)
        calcbtn.pack(side=LEFT,padx=(1,50),pady=(0,5),ipady=(5))


#==========================================================| |  Search  | |===========================================

        refreshFrame = Frame(master=fram_heshov2 ,relief="solid",border=0)
        refreshFrame.config(bg=primary)
        refreshFrame.pack(side='left',padx=0,ipady=0)

        # | כפתור הצג הכל| 
        kaftor(text='הצג הכל',master=refreshFrame,command=get_api)




# ------------------- | בחירת שם חיפוש | -----------------------------------
        laBel(fram_heshov2,'חיפוש',side=RIGHT)

        search = cttk.Combobox(fram_heshov2,style="Custom.TCombobox", state='readonly',width=15, justify='center')
        search.pack(side='right', pady=(0,5),padx=(0,30))
        search.configure(bootstyle='primary')
        search.bind("<<ComboboxSelected>>", Search_function)


        def get_workers_names():
            try:
                connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                cursor = connection.cursor()

                cursor.execute("SELECT DISTINCT workername FROM workers ORDER BY workername ASC;")

                rows = cursor.fetchall()

                global worker_names
                worker_names = [row[0] for row in rows]
            
                search['values'] = worker_names

            except:
                pass
        get_workers_names()
#=====================================================================================================================



        f7 = Frame(self.new_root, bd=0)
        f7.pack(pady=2, fill='both', expand=True)


        
#============================= | |  Main Content Area   | | ====================================

        #------------------- |  מחיקה | ---------

        def mainContent():
            deleteframe = Frame(f7,)
            deleteframe.config(bg=primary)
            deleteframe.pack(side=LEFT,fill='y')

            laBel(deleteframe,"מחק לפי מ'ס שורה")

            En_Delete = cttk.Entry(deleteframe,
                                    textvariable=delete_var,
                                    justify='center',
                                    font=(mainFont ,10),
                                    bootstyle='danger')
            En_Delete.pack(pady=0)


            #==  כפתור מחיקה |=

            kaftor('מחיקה',DeleteAndClear_function,deleteframe,side='bottom')
    #=====================================================================================================================





    #======================== | שדות כלט | =============================
            f1 = cttk.Frame(f7,bootstyle='dark',relief='flat',border=5)
            f1.pack(side='right',fill='y',pady=0)

            def en_Title():
                global entry_and_entryTitle
                global En1
                def entry_and_entryTitle(master,validcommand, valid,textvar, placehold):
                    En1 = cttk.Entry(master=master,bootstyle='default',takefocus=placehold,validatecommand=validcommand,validate=valid, textvariable=textvar, font=(mainFont, 10,'bold'), justify='center')
                    En1.insert(0, placehold)
                    En1.bind('<Enter>', lambda e: on_entry_focus_in(En1, placehold))
                    En1.bind('<Leave>', lambda e: on_entry_focus_out(En1, placehold))
                    En1.pack(pady=(3,0))

                    def on_entry_focus_in(entry, placeholder):
                        if entry.get() == placeholder:
                            entry.delete(0, 'end')

                    def on_entry_focus_out(entry, placeholder):
                        if not entry.get():
                            entry.insert(0, placeholder)

                # register the validation callback
                global alpha_func
                digit_func = self.new_root.register(validate_number)
                alpha_func = self.new_root.register(validate_alpha)

                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_NAME_var ,placehold = 'שם')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_PHONE_var ,placehold = 'טלפון')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_ID_var ,placehold= 'ת.ז')
                def get_date_from_entry():
                    global date_entry
                    date_entry = cttk.DateEntry(f1, dateformat='%x', firstweekday=1, startdate=None, style='default').pack(pady=(3, 0))
                def update_date_var():
                    entered_date = date_entry.get()
                    WORKERS_DATE_var.set(entered_date)
                    print(entered_date)


                # Create DateEntry widget and bind update function
                update_date_var = get_date_from_entry()





                # entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_DATE_var ,placehold = '')
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_COMPANY_NAME_var ,placehold = 'שם חברה')
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_COMPANY_MANAGER_NAME_var ,placehold = ' מנהל עבודה') 
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_WORK_ADDRESS_var ,placehold = 'כתובת')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_WAGE_var ,placehold = 'שח')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_Taken_var ,placehold = '0')
                # kaftor('hhh',update_date_var,f1)

                
                #========================= | כפתור הןספה | ======================
                button1 = cttk.Button(f1, text='הוספת', command=addandclear_function, cursor='hand2 ')
                button1.pack(side=BOTTOM,pady=(0,10))
                ToolTip(button1,delay=250, text="מלא את שדות הכלט מעלה לפני ההוספה", bootstyle=('dark', INVERSE))
            en_Title()

    #=======================================================================



# ===================== | |  Buttons frame | | =======================================

# ===================== | פרים כפתורים | =======================================
            bottonFrame = Frame(f7,bd=3,relief='flat',bg='gray')
            bottonFrame.pack(side='right',anchor='nw',fill='y')


            entry_and_entryTitle(master=bottonFrame,validcommand=(alpha_func, '%P'),valid='focus',textvar=WORKERS_Hours_Var,placehold=str("מ'ס שעות"))

            kaftor(text='תיקון',command=updateandclear,master=bottonFrame)

            kaftor(text="Docx הדפסה ל",command=print_to_docx,master=bottonFrame)

            kaftor(text="Exel הדפסה ל",command=print_to_excel,master=bottonFrame)

            kaftor(text='צור קשר',command=infoo,master=bottonFrame)

            kaftor(text='נקה שדות כלט',command=Clear_function,master=bottonFrame)

            kaftor(text='סגור תוכנה',command=des,master=bottonFrame)
#======================================================================================


            # | פרים Rowid | -----------------------------
            rowIdFrame = Frame(f7,bg='blue',relief='sunken',bd=3)
            rowIdFrame.pack(fill='both',ipady=5)
            #----------------------------------------------------------


            # ====================== | ניתוח נתונים |
            
            # | פרים ניתוח נתונים | -----------------------------
            underprintFrame = cttk.Frame (f7)
            underprintFrame.pack(pady=0,fill='both',ipady=300)
            

            global dataAnalysis_func
            def dataAnalysis_func():
                global get_workers_names2
                calc_Days_Frame = cttk.Frame(underprintFrame)
                calc_Days_Frame.pack(fill='both',ipady=300)

                # =========== | לאביל ניתוח נתונים | =============
                laBel(calc_Days_Frame,'ניתוח נתונים')
                global search_ttk
                search_ttk = cttk.Combobox(master=calc_Days_Frame,
                                        style="Custom.TCombobox",
                                        validate='all',
                                        state='readonly',
                                        font=(mainFont,10),
                                        justify='center')
                search_ttk.pack(side='top',pady=(10, 20))

                global get_workers_names_Analysis
                def get_workers_names_Analysis():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        cursor = connection.cursor()

                        cursor.execute("SELECT DISTINCT workername FROM workers;")

                        rows = cursor.fetchall()

                        worker_names = [row[0] for row in rows]

                        search_ttk['values'] = [''] + worker_names

                    except pymysql.err.DatabaseError as e:
                        toastErrorCacher("no Interner Error",f"אירעה שגיאה בעת חיבור למסד הנתונים {e}")
                    finally:
                        pass
                get_workers_names_Analysis()

                def howMutchDays():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        
                        cursor = connection.cursor()
                        cursor.execute(f"SELECT COUNT(*) AS name_count FROM workers WHERE workername = '{search_ttk.get()}'")
                        result = cursor.fetchone()
                        
                        if result:
                            name_count = result[0]
                            fresh = '      '
                            messagebox.showinfo("",f"מספר הימים ל {search_ttk.get()}\n{fresh}   {name_count} ")
                                
                        else:
                            messagebox.showinfo(f"ERROR","אין שורות המכילות את שם העובד")
                            
                    except pymysql.err.DatabaseError as e:
                        toastErrorCacher("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                    finally:
                        pass

                def takenMoney():
                    try:

                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        
                        cursor = connection.cursor()

                        cursor.execute(f"SELECT workername, SUM(`Taken`) AS total_wage FROM workers WHERE workername = '{search_ttk.get()}'")
                        
                        result = cursor.fetchone()
                        
                        if result:
                            name_count = result[1]
                            fresh = '      '
                            messagebox.showinfo("", f"   קיבל מפריעה  \n{fresh}   {name_count} ")
                                
                        else:
                            messagebox.showinfo(f"ERROR", "אין שורות המכילות את שם העובד")
                            
                    except pymysql.err.DatabaseError as e:
                        # Handling database errors
                        toastErrorCacher("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                    finally:
                        pass

                def howMutchDaysAll():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        cursor = connection.cursor()
                        cursor.execute("SELECT COUNT(DISTINCT workername) AS name_count FROM workers")
                        result = cursor.fetchone()
                        
                        if result:
                            name_count = result[0]
                            messagebox.showinfo("מספר העובדים:",f"מספר העובדים הוא: {search_ttk.get()}{name_count} ")
                                
                        else:
                            messagebox.showinfo(f"ERROR","אין שורות המכילות את שם העובד")
                            
                    except pymysql.err.DatabaseError as e:
                        toastErrorCacher("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                        pass


                def Company():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta,  database=database1)
                        cursor = connection.cursor()
                        
                        selected_value = search_ttk.get()

                        query = f"""
                                SELECT compname, workername, COUNT(*) AS total_days_worked 
                                FROM workers 
                                WHERE workername = '{selected_value}' 
                                GROUP BY compname, workername,hours
                            """

                        cursor.execute(query)
                        results = cursor.fetchall()
                        
                        if results:
                            global messageFormCompany
                            messageFormCompany = "פרטי החברות\n\n \n\nמספר ימי עבודה בכל חברה\n\n"
                            for result in results:
                                company_name = result[0]
                                workername = result[1]
                                days_worked = result[2]
                                messageFormCompany += f"_______________________________________________\n\nשם חברה: {company_name}\n\n מזהה עובד: {workername}\n\nימים שעבד: {days_worked}\n\n"
                            cttk.dialogs.dialogs.Messagebox.ok(messageFormCompany, title=' ', alert=False, parent=None)
                            # messagebox.showinfo("פרטי העבודה בחברות", message)
                        else:
                            messagebox.showinfo("ERROR", "אין נתונים להצגה")
                    except pymysql.err.DatabaseError as e:
                        toastErrorCacher("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                        return None

                def hours():
                    try:
                        conn = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                        cursor = conn.cursor()
                        worker_name = search_ttk.get()
                        query = f"SELECT workername, SUM(`hours`) AS total_wage FROM workers WHERE workername = '{worker_name}' GROUP BY workername"
                        cursor.execute(query)
                        res = cursor.fetchall()
                        total_wage = res[0]
                        if total_wage:
                            messagebox.showinfo('שעות עבודה', f'ס"כ שעות עבודה: {total_wage}')
                        else:
                            messagebox.showinfo('אין תוצאות', 'לא נמצאו נתונים עבור העובד המבוקש')
                    except Exception as e:
                        toastErrorCacher('לא ניתן להתחבר לשרת MySQL',
                                        f"לא ניתן היה ליצור חיבור מכיוון שמכונת היעד סירבה לכך באופן פעיל. בדוק את חיבור האינטרנט שלך ונסה שנית. השגיאה היא: {e}", 'danger')
                    finally:
                        pass

                buttonsframe = Frame(master=calc_Days_Frame,relief='groove')
                buttonsframe.config(bg=dark)
                buttonsframe.pack(side=RIGHT,fill='x')


                def generate_excel():
                    search_value = search_ttk.get()
                    try:
                        # Connect to MySQL database
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        cursor = connection.cursor()

                        # Define the SQL query with parameter placeholder %s
                        query = """
                            SELECT compname, workername, COUNT(*) AS total_days_worked 
                            FROM workers 
                            WHERE workername = %s
                            GROUP BY compname, workername
                        """

                        cursor.execute(query, (search_value,))
                        results = cursor.fetchall()

                        # If there are results, create Excel file
                        if results:
                            # Create a new Excel workbook
                            excel_filename = f"{search_value}.xlsx"
                            excel_path = os.path.join(os.getcwd(), excel_filename)
                            workbook = xlsxwriter.Workbook(excel_path)
                            worksheet = workbook.add_worksheet("דוח חברה")

                            # Define formats for header and data
                            header_format = workbook.add_format({'bold': True, 'align': 'right', 'text_wrap': True})
                            data_format = workbook.add_format({'align': 'right', 'text_wrap': True})

                            # Write headers
                            headers = ["שם חברה", "מזהה עובד", "ימים שעבד"]
                            for col, header in enumerate(headers):
                                worksheet.write(0, col, header, header_format)

                            # Write data rows
                            for row_idx, result in enumerate(results, start=1):
                                company_name = result[0]
                                worker_name = result[1]
                                days_worked = result[2]
                                worksheet.write(row_idx, 0, company_name, data_format)
                                worksheet.write(row_idx, 1, worker_name, data_format)
                                worksheet.write(row_idx, 2, days_worked, data_format)

                            # Adjust column widths based on content length
                            for col_idx, header in enumerate(headers):
                                max_length = max(len(str(header)), max(len(str(result[col_idx])) for result in results))
                                worksheet.set_column(col_idx, col_idx, max_length + 4)

                            # Close the workbook
                            workbook.close()

                            # Show success message
                            messagebox.showinfo("נוצר קובץ Excel", f"נוצר קובץ Excel בשם: {excel_filename} בתיקייה: {excel_path}")

                        else:
                            # Show message if no results found
                            messagebox.showinfo("אין נתונים", "לא נמצאו נתונים עבור העובד המבוקש")

                    except pymysql.err.DatabaseError as e:
                        # Handle database errors
                        print(e)
                        messagebox.showerror("שגיאה בגישה למסד הנתונים", f"אירעה שגיאה בגישה למסד הנתונים: {str(e)}")

                    except Exception as e:
                        # Handle other errors
                        print(e)
                        messagebox.showerror("שגיאה ביצירת קובץ Excel", f"אירעה שגיאה ביצירת קובץ Excel: {str(e)}")

                    finally:
                        # Close database connection
                        if connection:
                            connection.close()


        # ====================== | כפתורי בדיקה |
                def dataAnalysis(text,command):
                    btn2 = cttk.Button(master=buttonsframe,
                                        text=text,
                                        command=command,
                                        width=30
                                        );
                    btn2.pack(pady=2,fill=X)
                
                dataAnalysis(text='ימי עבודה ',command=howMutchDays)
                dataAnalysis(text='קיבל מפריעה',command=takenMoney)
                dataAnalysis(text='כמה עובדים יש',command=howMutchDaysAll)
                dataAnalysis(text='חברה',command=Company)
                dataAnalysis(text='שעות עבודה',command=hours)
                dataAnalysis(text='הדפסה',command=generate_excel)
                # dataAnalysis(text='שעות עבודה',command=generate_pdf(results))
            dataAnalysis_func()
    #----------------------------------------------------------

            

            # | מס שורה |
            laBel(rowIdFrame,'שורה')

    #--------------------------------------------------

            # | שדה מס שורה  |
            iidentry = cttk.Entry(rowIdFrame,background=dark,foreground=danger,font=(mainFont,10,'bold'),textvariable=Row_Number_var,state='disabled',justify='center')
            iidentry.pack(pady=0,side='bottom')
        mainContent()

        self.new_root.mainloop()
if __name__ == '__main__':
    MainScreen()
